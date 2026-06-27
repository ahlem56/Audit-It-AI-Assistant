from __future__ import annotations

import re
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate, RichText

from app.models.export_models import ExportReportRequest


# Static assets must stay outside ``app/data`` because Docker Compose mounts a
# persistent volume over that directory, which would otherwise hide templates
# baked into a newly rebuilt image.
WORD_TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "templates" / "rapport_template_dynamic.docx"

PWC_ORANGE = "D04A02"
PWC_RED = "E0301E"
PWC_DARK = "2D2D2D"
PWC_PEACH = "F7DED3"
PWC_LIGHT_GREY = "F2F2F2"

WORD_SECTIONS = [
    "Préambule",
    "Modalité et intervenants",
    "Objectifs et approche",
    "Périmètre d'intervention",
    "Synthèse générale",
    "Points relevés",
    "Plan d'action et recommandations détaillées",
    "Annexes",
    "Conclusion",
]


def _safe(value: object) -> str:
    return "" if value is None else str(value)


def _display_priority(priority: str) -> str:
    return {
        "Critical": "Critique",
        "High": "Élevée",
        "Medium": "Moyenne",
        "Low": "Faible",
    }.get(priority, priority or "")


def _extract_report_year(data) -> str:
    title = _safe(getattr(data, "cover_title", ""))
    match = re.search(r"FY\s*([0-9]{2,4})", title, flags=re.IGNORECASE)
    if match:
        year = match.group(1)
        return f"20{year}" if len(year) == 2 else year

    report_date = _safe(getattr(data, "report_date", ""))
    digits = "".join(char for char in report_date if char.isdigit())
    if len(digits) >= 4:
        return digits[-4:]

    period = _safe(getattr(data, "report_period", ""))
    tokens = [token for token in re.split(r"\D+", period) if len(token) == 4]
    return tokens[-1] if tokens else ""


def _stakeholder_item(value: object) -> dict[str, str]:
    if isinstance(value, dict):
        return {
            "name": _safe(value.get("name") or value.get("interlocuteur") or value.get("label")),
            "role": _safe(value.get("role") or value.get("fonction") or value.get("function")),
        }

    text = _safe(value).strip()
    role_match = re.match(r"^(?P<name>.+?)\s*\((?P<role>[^)]+)\)\s*$", text)
    if role_match:
        return {"name": role_match.group("name").strip(), "role": role_match.group("role").strip()}
    if " - " in text:
        name, role = text.split(" - ", 1)
        return {"name": name.strip(), "role": role.strip()}
    if ":" in text:
        name, role = text.split(":", 1)
        return {"name": name.strip(), "role": role.strip()}
    return {"name": text, "role": ""}


def _stakeholder_items(values: list[object]) -> list[dict[str, str]]:
    items: list[dict[str, str]] = []
    for value in values:
        if isinstance(value, dict):
            items.append(_stakeholder_item(value))
            continue

        text = _safe(value).strip()
        split_text = re.sub(r"\)\s+(?=[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ' -]+\s+\()", ")|", text)
        for part in [chunk.strip() for chunk in split_text.split("|") if chunk.strip()]:
            items.append(_stakeholder_item(part))
    return items


def _priority_summary_item(item) -> dict[str, Any]:
    percentage = getattr(item, "percentage", 0)
    try:
        percentage_value: int | float = round(float(percentage), 1)
    except (TypeError, ValueError):
        percentage_value = 0

    return {
        "priority": _safe(getattr(item, "priority", "")),
        "priority_label": _display_priority(_safe(getattr(item, "priority", ""))),
        "count": getattr(item, "count", 0),
        "percentage": percentage_value,
    }


def _priority_count(summary: list[dict[str, Any]], priority: str) -> int:
    for item in summary:
        if item.get("priority") == priority:
            try:
                return int(item.get("count") or 0)
            except (TypeError, ValueError):
                return 0
    return 0


def _priority_percentage(summary: list[dict[str, Any]], priority: str) -> str:
    for item in summary:
        if item.get("priority") == priority:
            try:
                return f"{round(float(item.get('percentage') or 0), 1):g}%"
            except (TypeError, ValueError):
                return "0%"
    return "0%"


def _priority_percentage_value(summary: list[dict[str, Any]], priority: str) -> float:
    for item in summary:
        if item.get("priority") == priority:
            try:
                return float(item.get("percentage") or 0)
            except (TypeError, ValueError):
                return 0.0
    return 0.0


def _compact_text(value: object, limit: int) -> str:
    text = re.sub(r"\s+", " ", _safe(value)).strip()
    if len(text) <= limit:
        return text
    shortened = text[: limit + 1].rsplit(" ", 1)[0].rstrip(" ,;:-")
    return f"{shortened}."


def _reference_findings_text(findings: list[dict[str, Any]], *, limit: int = 6) -> str:
    blocks: list[str] = []
    for index, finding in enumerate(findings[:limit], start=1):
        priority = finding.get("priority_label") or finding.get("priority") or "À qualifier"
        title = _compact_text(finding.get("title"), 90)
        observation = _compact_text(finding.get("finding"), 260)
        risk = _compact_text(finding.get("risk_impact"), 180)
        application = _compact_text(finding.get("application"), 45)
        blocks.append(
            f"{index:02d}  {title}  |  {priority}\n"
            f"{application} — {observation}\n"
            f"Risque : {risk}"
        )
    return "\n\n".join(blocks) or "Aucun constat reportable identifié sur le périmètre revu."


def _reference_recommendations_text(recommendations: list[dict[str, Any]], *, limit: int = 4) -> str:
    blocks: list[str] = []
    for index, item in enumerate(recommendations[:limit], start=1):
        title = _compact_text(item.get("title"), 80)
        action = _compact_text(item.get("recommendation"), 310)
        owner = _compact_text(item.get("owner") or item.get("owners"), 60) or "Responsable à confirmer"
        blocks.append(f"{index:02d}  {title}\n{action}\nResponsable : {owner}")
    return "\n\n".join(blocks) or "Aucune recommandation détaillée à présenter."


def _rich_text_block(text: str, *, size: int = 18) -> RichText:
    return RichText(text, font="Arial", size=size, color="2D2D2D")


def _bar(value: float, maximum: float = 100.0, width: int = 18) -> str:
    filled = 0 if maximum <= 0 else round(max(0.0, min(value, maximum)) / maximum * width)
    return "█" * filled + "░" * (width - filled)


def _cover_lines(data) -> tuple[str, str]:
    title = _safe(getattr(data, "cover_title", "")).strip()
    client = _safe(getattr(data, "client_name", "")).strip() or "Client"
    parts = [part.strip() for part in re.split(r"\s*[–—]\s*", title) if part.strip()]
    if parts and parts[0].lower().startswith(("audit itgc", "rapport d'audit")):
        parts = parts[1:]
    if len(parts) >= 2:
        return parts[0], " — ".join(parts[1:])

    client_parts = [part.strip() for part in re.split(r"\s*[–—]\s*", client) if part.strip()]
    display_client = next((part for part in reversed(client_parts) if "banque" in part.lower()), client_parts[-1] if client_parts else client)
    subject = title or _safe(getattr(data, "cover_subtitle", "")) or "Revue des contrôles généraux informatiques"
    return display_client, subject


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margin(cell, *, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_run(run, *, size: float = 9.5, bold: bool = False, color: str = PWC_DARK) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _add_banner(document, number: str, title: str, subtitle: str = "") -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(14.3)
    left, right = table.rows[0].cells
    left.width = Cm(2.0)
    right.width = Cm(14.3)
    _shade_cell(left, PWC_ORANGE)
    _shade_cell(right, PWC_PEACH)
    for cell in (left, right):
        _set_cell_margin(cell, top=180, bottom=180, start=180, end=180)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left_paragraph = left.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_run(left_paragraph.add_run(number), size=22, bold=True, color="FFFFFF")
    right_paragraph = right.paragraphs[0]
    _format_run(right_paragraph.add_run(title), size=16, bold=True)
    if subtitle:
        sub = right.add_paragraph()
        _format_run(sub.add_run(subtitle), size=8.5, color="5A5A5A")
    document.add_paragraph()


def _add_content_block(document, heading: str, content: object, *, fill: str | None = None) -> None:
    text = _safe(content).strip()
    if not text:
        return
    table = document.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header, body = table.rows[0].cells[0], table.rows[1].cells[0]
    _shade_cell(header, PWC_DARK)
    _shade_cell(body, fill or "FFFFFF")
    _set_cell_margin(header, top=80, bottom=80)
    _set_cell_margin(body, top=130, bottom=130)
    _format_run(header.paragraphs[0].add_run(heading), size=9, bold=True, color="FFFFFF")
    body_paragraph = body.paragraphs[0]
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_paragraph.paragraph_format.line_spacing = 1.08
    _format_run(body_paragraph.add_run(text), size=9.2)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _add_metadata_table(document, finding) -> None:
    priority = _safe(getattr(finding, "priority", ""))
    priority_fill = {"Critical": "C00000", "High": PWC_RED, "Medium": "FFB600", "Low": "70AD47"}.get(priority, PWC_DARK)
    values = (
        ("Référence", _safe(getattr(finding, "reference", ""))),
        ("Domaine", _safe(getattr(finding, "domain", ""))),
        ("Application", _safe(getattr(finding, "application", ""))),
        ("Priorité", _display_priority(priority)),
        ("Responsable", _safe(getattr(finding, "owner", "")) or _safe(getattr(finding, "owners", "")) or "À confirmer"),
        ("Catégorie", _safe(getattr(finding, "category", "")) or "ITGC"),
    )
    table = document.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index in range(3):
        for pair_index in range(2):
            label, value = values[row_index * 2 + pair_index]
            label_cell = table.cell(row_index, pair_index * 2)
            value_cell = table.cell(row_index, pair_index * 2 + 1)
            _shade_cell(label_cell, PWC_LIGHT_GREY)
            if label == "Priorité":
                _shade_cell(value_cell, priority_fill)
            for cell in (label_cell, value_cell):
                _set_cell_margin(cell, top=75, bottom=75)
            _format_run(label_cell.paragraphs[0].add_run(label), size=8.2, bold=True)
            _format_run(
                value_cell.paragraphs[0].add_run(value),
                size=8.4,
                bold=label == "Priorité",
                color="FFFFFF" if label == "Priorité" else PWC_DARK,
            )
    document.add_paragraph()


def _append_page_number_footer(document, client_name: str, report_year: str) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[-1] if section.footer.paragraphs else section.footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_before = Pt(4)
        run = paragraph.add_run(f"pwc  |  Rapport d'audit ITGC — {client_name} — {report_year}  |  ")
        _format_run(run, size=7.5, color="666666")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, instruction, separate, end))


def _add_compact_finding_card(document, finding, index: int) -> None:
    priority = _safe(getattr(finding, "priority", ""))
    priority_fill = {"Critical": "C00000", "High": PWC_RED, "Medium": "FFB600", "Low": "70AD47"}.get(priority, PWC_DARK)
    table = document.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    title_cell = table.cell(0, 0).merge(table.cell(0, 1))
    _shade_cell(title_cell, priority_fill)
    _set_cell_margin(title_cell, top=85, bottom=85, start=120, end=120)
    title_paragraph = title_cell.paragraphs[0]
    _format_run(
        title_paragraph.add_run(f"{index:02d}  {_safe(getattr(finding, 'title', ''))}"),
        size=10.2,
        bold=True,
        color="FFFFFF",
    )
    priority_run = title_paragraph.add_run(f"    {_display_priority(priority).upper()}")
    _format_run(priority_run, size=8.2, bold=True, color="FFFFFF")

    metadata_cell = table.cell(1, 0).merge(table.cell(1, 1))
    _shade_cell(metadata_cell, PWC_LIGHT_GREY)
    _set_cell_margin(metadata_cell, top=55, bottom=55, start=120, end=120)
    metadata = "  |  ".join(
        part
        for part in (
            f"Réf. {_safe(getattr(finding, 'reference', ''))}",
            _safe(getattr(finding, "domain", "")),
            _safe(getattr(finding, "application", "")),
            f"Responsable : {_safe(getattr(finding, 'owner', '')) or _safe(getattr(finding, 'owners', '')) or 'À confirmer'}",
        )
        if part
    )
    _format_run(metadata_cell.paragraphs[0].add_run(metadata), size=7.8, bold=True, color="555555")

    row_content = (
        ("Contrôle attendu", getattr(finding, "expected_control", ""), "FAFAFA"),
        ("Constat", getattr(finding, "finding", ""), "FFFFFF"),
        ("Risque / cause", "\n".join(filter(None, (_safe(getattr(finding, "risk_impact", "")), _safe(getattr(finding, "root_cause", ""))))), "FFF4F1"),
        ("Recommandation", getattr(finding, "recommendation", ""), "FFF8E8"),
    )
    for row_index, (label, value, fill) in enumerate(row_content, start=2):
        label_cell, value_cell = table.rows[row_index].cells
        label_cell.width = Cm(3.0)
        value_cell.width = Cm(13.3)
        _shade_cell(label_cell, PWC_DARK)
        _shade_cell(value_cell, fill)
        _set_cell_margin(label_cell, top=55, bottom=55, start=90, end=90)
        _set_cell_margin(value_cell, top=55, bottom=55, start=100, end=100)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _format_run(label_cell.paragraphs[0].add_run(label), size=7.7, bold=True, color="FFFFFF")
        paragraph = value_cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        _format_run(paragraph.add_run(_safe(value)), size=7.8)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    spacer.paragraph_format.space_before = Pt(0)


def _append_complete_findings(document, data, report_year: str) -> None:
    findings = list(data.detailed_findings or [])
    document.add_page_break()
    _add_banner(
        document,
        "07",
        "Fiches détaillées des constats",
        f"{len(findings)} constat(s) — périmètre complet de la mission",
    )
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    _format_run(
        intro.add_run(
            "Cette section reprend l'intégralité des constats, risques, causes et recommandations afin de garantir un rapport auditable et exploitable."
        ),
        size=9.5,
    )

    for index, finding in enumerate(findings, start=1):
        if index > 1 and index % 2 == 1:
            document.add_page_break()
        _add_compact_finding_card(document, finding, index)

    document.add_page_break()
    _add_banner(document, "08", "Conclusion", f"Rapport d'audit ITGC — {report_year}")
    _add_content_block(document, "Conclusion générale", data.conclusion or data.general_synthesis, fill=PWC_PEACH)
    _append_page_number_footer(document, _safe(data.client_name) or "Client", report_year)


def _finding_item(finding) -> dict[str, Any]:
    priority = _safe(getattr(finding, "priority", ""))
    recommendation_steps = getattr(finding, "recommendation_steps", []) or []
    action_text = _safe(getattr(finding, "recommendation", ""))
    if recommendation_steps:
        action_text = action_text + "\n" + "\n".join(f"- {step}" for step in recommendation_steps)

    return {
        "observation_id": _safe(getattr(finding, "observation_id", "")),
        "reference": _safe(getattr(finding, "reference", "")),
        "domain": _safe(getattr(finding, "domain", "")),
        "category": _safe(getattr(finding, "category", "")),
        "application": _safe(getattr(finding, "application", "")),
        "layer": _safe(getattr(finding, "layer", "")),
        "owners": _safe(getattr(finding, "owners", "")),
        "title": _safe(getattr(finding, "title", "")),
        "expected_control": _safe(getattr(finding, "expected_control", "")),
        "finding": _safe(getattr(finding, "finding", "")),
        "risk_impact": _safe(getattr(finding, "risk_impact", "")),
        "risk_scenario": _safe(getattr(finding, "risk_scenario", "")),
        "impact_detail": _safe(getattr(finding, "impact_detail", "")),
        "business_impact": _safe(getattr(finding, "business_impact", "")),
        "control_impact": _safe(getattr(finding, "control_impact", "")),
        "compliance_impact": _safe(getattr(finding, "compliance_impact", "")),
        "root_cause": _safe(getattr(finding, "root_cause", "")),
        "aggravating_factors": list(getattr(finding, "aggravating_factors", []) or []),
        "immediate_action": _safe(getattr(finding, "immediate_action", "")),
        "structural_action": _safe(getattr(finding, "structural_action", "")),
        "owner": _safe(getattr(finding, "owner", "")),
        "evidence_expected": _safe(getattr(finding, "evidence_expected", "")),
        "follow_up_mechanism": _safe(getattr(finding, "follow_up_mechanism", "")),
        "recommendation": action_text,
        "auditor_comment": _safe(getattr(finding, "auditor_comment", "")),
        "management_summary": _safe(getattr(finding, "management_summary", "")),
        "priority": priority,
        "priority_label": _display_priority(priority),
    }


def _application_detail_item(application) -> dict[str, str]:
    return {
        "name": _safe(getattr(application, "name", "")),
        "description": _safe(getattr(application, "description", "")),
        "operating_system": _safe(getattr(application, "operating_system", "")),
        "database": _safe(getattr(application, "database", "")),
        "provider": _safe(getattr(application, "provider", "")),
    }


def _word_context(result: ExportReportRequest) -> dict[str, Any]:
    data = result.structured_output
    report_year = _extract_report_year(data)
    previous_year = str(int(report_year) - 1) if report_year.isdigit() else ""
    priority_summary = [_priority_summary_item(item) for item in (data.priority_summary or [])]
    findings = [_finding_item(item) for item in (data.detailed_findings or [])]
    recommendations = [_finding_item(item) for item in (data.detailed_recommendations or [])]
    applications = list(data.applications or [])
    application_details = [_application_detail_item(item) for item in (data.application_details or [])]
    if not application_details:
        application_details = [
            {
                "name": _safe(application),
                "description": "",
                "operating_system": "",
                "database": "",
                "provider": "",
            }
            for application in applications
        ]
    maturity_level = _safe(getattr(data, "maturity_level", "")) or "À confirmer"
    critical_count = _priority_count(priority_summary, "Critical")
    high_count = _priority_count(priority_summary, "High")
    total_findings = len(findings)
    domain_counts = Counter(item.get("domain") or "Autres" for item in findings)
    domain_distribution = "\n".join(
        f"{domain} : {count}" for domain, count in domain_counts.most_common(6)
    ) or "Aucune exposition significative identifiée."
    distribution = " | ".join(
        (
            f"Critique {_priority_percentage(priority_summary, 'Critical')}",
            f"Élevée {_priority_percentage(priority_summary, 'High')}",
            f"Moyenne {_priority_percentage(priority_summary, 'Medium')}",
            f"Faible {_priority_percentage(priority_summary, 'Low')}",
        )
    )
    priority_chart = "\n".join(
        f"{label:<9} {_bar(_priority_percentage_value(priority_summary, key))}  {_priority_percentage(priority_summary, key)}"
        for key, label in (("Critical", "Critique"), ("High", "Élevée"), ("Medium", "Moyenne"), ("Low", "Faible"))
    )
    max_domain_count = max(domain_counts.values(), default=0)
    domain_chart = "\n".join(
        f"{domain}\n{_bar(float(count), float(max_domain_count or 1), width=14)}  {count} constat(s)"
        for domain, count in domain_counts.most_common(5)
    ) or "Aucune exposition significative identifiée."
    general_synthesis = _safe(data.general_synthesis) or _safe(data.executive_summary)
    scope_analysis = "\n\n".join(
        part
        for part in (
            _compact_text(data.scope_summary, 750),
            "Applications : " + ", ".join(applications) if applications else "",
            "Approche : " + " • ".join(data.audit_approach or []) if data.audit_approach else "",
        )
        if part
    )
    cover_title = _safe(data.cover_title) or "Rapport d'audit IT"
    client_name = _safe(data.client_name) or "Client"
    cover_client, cover_subject = _cover_lines(data)

    return {
        "client_name": client_name,
        "report_year": report_year,
        "prior_fiscal_year": f"FY{previous_year[-2:]}" if previous_year else "",
        "report_period": _safe(data.report_period),
        "report_date": _safe(data.report_date),
        "cover_title": cover_title,
        "cover_client": cover_client,
        "cover_subject": cover_subject,
        "cover_subtitle": _safe(data.cover_subtitle),
        "confidentiality_notice": _safe(data.confidentiality_notice) or "Strictement privé et confidentiel",
        "word_sections": WORD_SECTIONS,
        "table_of_contents": list(data.table_of_contents or WORD_SECTIONS),
        "preamble": _safe(data.preamble),
        "modality": _safe(data.scope_summary),
        "objectives": list(data.objectives or []),
        "audit_approach": list(data.audit_approach or []),
        "scope_summary": _safe(data.scope_summary),
        "applications": applications,
        "application_details": application_details,
        "stakeholders": _stakeholder_items(list(data.stakeholders or [])),
        "priority_summary": priority_summary,
        "metrics": {
            "total_findings": len(findings),
            "critical_count": _priority_count(priority_summary, "Critical"),
            "high_count": _priority_count(priority_summary, "High"),
            "applications_count": len(application_details) or len(applications),
            "maturity_level": maturity_level,
            "maturity_assessment": _safe(getattr(data, "maturity_assessment", "")) or "Appréciation issue de la synthèse des constats.",
        },
        "general_synthesis": _safe(data.general_synthesis) or _safe(data.executive_summary),
        "executive_highlights": list(getattr(data, "executive_highlights", []) or []),
        "watch_points": list(data.watch_points or []),
        "detailed_findings": findings,
        "detailed_recommendations": recommendations,
        "appendices": list(data.appendices or []),
        "conclusion": _safe(data.conclusion),
        # Fields used by the original fixed-layout PwC reference document.
        "date": _safe(data.report_date) or report_year,
        "resume_text": _compact_text(general_synthesis, 1900),
        "priority_findings_count": critical_count + high_count,
        "total_findings_count": total_findings,
        "applications_count": len(application_details) or len(applications),
        "maturity_level": maturity_level,
        "low_percentage": _priority_percentage(priority_summary, "Low"),
        "high_percentage": _priority_percentage(priority_summary, "High"),
        "critical_percentage": _priority_percentage(priority_summary, "Critical"),
        "pourcentage": _priority_percentage(priority_summary, "Critical") if critical_count else _priority_percentage(priority_summary, "High"),
        "remarque": _compact_text(getattr(data, "priority_insight", "") or getattr(data, "maturity_assessment", ""), 220),
        "analyse_text": _compact_text(scope_analysis, 1800),
        "reparation_transactions_graphes": distribution,
        "niveau_risque_graphes": f"{critical_count} critique(s) | {high_count} élevée(s)",
        "exposition_financière_graphes": domain_distribution,
        "priority_chart": _rich_text_block(priority_chart, size=19),
        "risk_level_chart": _rich_text_block(
            f"Critique  {_bar(float(critical_count), float(max(total_findings, 1)), width=12)}  {critical_count}\n"
            f"Élevée   {_bar(float(high_count), float(max(total_findings, 1)), width=12)}  {high_count}",
            size=20,
        ),
        "domain_chart": _rich_text_block(domain_chart, size=18),
        "transaction_text": _rich_text_block(_reference_findings_text(findings)),
        "critical_case_text": _rich_text_block(_reference_findings_text(
            [item for item in findings if item.get("priority") in {"Critical", "High"}], limit=2,
        )),
        "recommendations_text": _rich_text_block(_reference_recommendations_text(recommendations)),
        "glossary_text": _rich_text_block("\n".join(data.appendices or [
            "ITGC : contrôles généraux informatiques.",
            "Critique : action immédiate requise.",
            "Élevée : remédiation à court terme.",
            "Moyenne : amélioration planifiée.",
            "Faible : optimisation recommandée.",
        ])),
    }


def build_report_docx(result: ExportReportRequest) -> BytesIO:
    if not WORD_TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Word template not found: {WORD_TEMPLATE_PATH}")

    document = DocxTemplate(str(WORD_TEMPLATE_PATH))
    document.render(_word_context(result))

    rendered = BytesIO()
    document.save(rendered)
    rendered.seek(0)

    complete_document = WordDocument(rendered)
    _append_complete_findings(
        complete_document,
        result.structured_output,
        _extract_report_year(result.structured_output),
    )
    output = BytesIO()
    complete_document.save(output)
    output.seek(0)
    return output
